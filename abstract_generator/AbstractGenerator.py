# -*- coding: utf-8 -*-
"""
Generate abstract document (docx) file from table (xlsx)
by nebula

Dependency: pandas, xlrd, python-docx, pillow
"""
from PIL import Image
import pandas as pd
import docx
import math
import re
import os


class AbstractGenerator:
    def __init__(self, image_dir='', template_type='aini2016'):
        self.records = None
        self.image_dir = image_dir
        self.template_type = template_type
        self.exreg4author = re.compile(r'^([^\)]+)((?:\(.+\))*)$')
        self.exreg4affiliation = re.compile(r'^((?:\(.+\))*)(.+)$')
        self.exreg4super = re.compile(r'(\(\w+\))')
        # self.exreg4italic = re.compile(r'(\<i\>\w+\</i\>)')
        self.exreg4italic = re.compile(r'(\<i\>.*?\</i\>)')
        self.exreg4sup = re.compile(r'(\<sup\>.*?\</sup\>)')
        self.exreg4sub = re.compile(r'(\<sub\>.*?\</sub\>)')
        self.exreg4tags = re.compile(r'(\<.*?\>.*?\</.*?\>)')
        self.exreg4tag_strip = re.compile(r'<[^>]*?>')
        self.preferredImageMaxWidth = 14  # cm
        self.preferredImageMaxHeight = 8.5  # cm
        self.preferredImageDpi = 72

    def _insert_image(self, filename, image_filename):
        doc = docx.Document(filename)

        for paragraph in doc.paragraphs:
            if '[[FIGURE]]' in paragraph.text:
                # paragraph.text = ''
                run = paragraph.add_run()
                run.add_paragraph()
                inline_shape = run.add_picture(image_filename, width=docx.shared.Pt(300))
                run.add_paragraph()

        doc.save(filename)

    @staticmethod
    def _empty(text):
        if isinstance(text, float) and math.isnan(text):
            return True
        return text.strip() == ''

    def _to_array(self, text, delim):
        if self._empty(text):
            return []
        items = text.split(delim)
        return [item for item in items if item.strip()]

    @staticmethod
    def _remove_parentheses(text):
        exreg = re.compile(r'\((\w+)\)')
        nums = exreg.split(text)
        num = ''
        for n in nums:
            n = n.strip()
            if n == '':
                continue
            if num != '':
                num += ', '
            num += n
        return num

    @staticmethod
    def _get_image_size(pixel, dpi):
        return pixel / dpi * 2.54

    def _get_preferred_image_size(self, fpath):
        img = Image.open(fpath)
        dpi = (self.preferredImageDpi, self.preferredImageDpi)
        if 'dpi' in img.info:
            dpi = img.info['dpi']
        if 'jfif_density' in img.info:
            dpi = img.info['jfif_density']
        width = self._get_image_size(img.size[0], dpi[0])
        height = self._get_image_size(img.size[1], dpi[1])
        if width > self.preferredImageMaxWidth:
            height = height * self.preferredImageMaxWidth / width
            width = self.preferredImageMaxWidth
        if height > self.preferredImageMaxHeight:
            width = width * self.preferredImageMaxHeight / height
            height = self.preferredImageMaxHeight
        # print('image: %s(w:%dpx(%gcm),h:%dpx(%gcm),dpi:%s) -> (w:%gcm,h:%gcm)' % (fpath, img.size[0], self._getImageSize(img.size[0], dpi[0]), img.size[1], self._getImageSize(img.size[1], dpi[1]), dpi, width, height))
        img.close()
        return docx.shared.Cm(width), docx.shared.Cm(height)

    def _apply_it_sup_sub(self, doc, body, debug=False):
        p = doc.add_paragraph()
        body_split = self.exreg4tags.split(body)

        for split in body_split:
            italic_mode = False
            sup_mode = False
            sub_mode = False

            if self.exreg4italic.match(split):
                italic_mode = True
            if self.exreg4sup.match(split):
                sup_mode = True
            if self.exreg4sub.match(split):
                sub_mode = True

            if debug:
                run = p.add_run(split)
            else:
                run = p.add_run(self.exreg4tag_strip.sub('', split))

            run.italic = italic_mode
            run.font.superscript = sup_mode
            run.font.subscript = sub_mode

        return p

    def read_xlsx(self, filename):
        print('Reading: %s' % filename)
        exls = pd.ExcelFile(filename)
        self.records = exls.parse()

    def write_docx(self, filename, template=None):
        print('Writing: %s' % filename)

        if template is not None:
            doc = docx.Document(template)
        else:
            doc = docx.Document()

        first = True
        for i in self.records.index:
            if first:
                section = doc.sections[0]
            else:
                section = doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
            section.orientation = docx.enum.section.WD_ORIENT.PORTRAIT
            section.page_height = docx.shared.Mm(297)
            section.page_width = docx.shared.Mm(210)
            section.top_margin = docx.shared.Mm(20)
            section.right_margin = docx.shared.Mm(20)
            section.left_margin = docx.shared.Mm(20)
            section.bottom_margin = docx.shared.Mm(15)
            if self.template_type == 'aini2016':
                self._write_doc_aini2016(doc, self.records.loc[i])
            else:
                self._write_doc_jscpb2016(doc, self.records.loc[i])
            first = False

        doc.save(filename)

    def _write_doc_jscpb2016(self, doc, record):
        print(record.title)

        # Title
        # p = doc.add_paragraph(record.title)
        p = self._apply_it_sup_sub(doc, record.title)
        for run in p.runs:
            run.font.size = docx.shared.Pt(12)
            run.bold = True

        # Authors
        p = doc.add_paragraph()
        author_list = self.exreg4super.split(record.authors)
        for j in range(len(author_list)):
            if j & 1:
                p.add_run(author_list[j]).font.superscript = True
            else:
                p.add_run(author_list[j])

        # Affiliations
        p = doc.add_paragraph(record.affiliations)
        p.runs[0].font.size = docx.shared.Pt(9)
        p.runs[0].italic = True

        # Abstract Body
        self._apply_it_sup_sub(doc, record.abstract)

        # keywords
        p = doc.add_paragraph('Keywords: ')
        p.add_run(record.keywords).italic = True

    def _write_doc_aini2016(self, doc, record):
        print('"%s"' % record['Title'])
        exreg4num = re.compile(r'\((\w+)\)')

        font = doc.styles['Normal'].font
        font.size = docx.shared.Pt(10)
        font.name = 'Times New Roman'

        # Program Number
        # p = doc.add_paragraph()
        # p.paragraph_format.line_spacing = docx.shared.Pt(12)
        # p.paragraph_format.space_after = docx.shared.Pt(5)
        # r = p.add_run(record['Program No.'].strip())

        # Title
        p = doc.add_paragraph()
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = docx.shared.Pt(25)
        p.paragraph_format.space_after = docx.shared.Pt(14)
        r = p.add_run(record['Title'].strip())
        r.font.size = docx.shared.Pt(12)
        r.bold = True
        r.italic = True

        # Authors
        p = doc.add_paragraph()
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = docx.shared.Pt(12)
        p.paragraph_format.space_after = docx.shared.Pt(12)
        authors = self._to_array(record['Name'], '\n')
        first = True
        for author in authors:
            m = self.exreg4author.match(author)
            if not first:
                p.add_run(', ').bold = True
            name = m.group(1).strip().replace(' ', '\u00A0')
            num = self._remove_parentheses(m.group(2).strip())
            p.add_run(name).bold = True
            if num != '':
                r = p.add_run('\u00A0' + num)
                r.bold = True
                r.font.superscript = True
            first = False
        p.add_run('\n')

        # Affiliation
        affiliations = self._to_array(record['Affiliation'], '\n')
        first = True
        for affiliation in affiliations:
            m = self.exreg4affiliation.match(affiliation)
            if not first:
                p.add_run(', ')
            num = self._remove_parentheses(m.group(1).strip())
            name = m.group(2).strip()
            if num != '':
                r = p.add_run(num + '\u00A0')
                r.font.superscript = True
            p.add_run(name)
            first = False
        p.add_run('\n' + record['e-mail'])

        # DOI
        p = doc.add_paragraph('DOI:' + record['DOI'].strip())
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = docx.shared.Pt(12)

        # Abstract Body
        items = self._to_array(record['Abstract'], '\n')
        first = True
        for item in items:
            p = doc.add_paragraph(item)
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = docx.shared.Pt(11)
            p.paragraph_format.space_after = docx.shared.Pt(2)
            if not first:
                p.paragraph_format.first_line_indent = docx.shared.Pt(12)
            first = False
        p.paragraph_format.space_after = docx.shared.Pt(12)

        # Figure
        if not self._empty(record['Figure file Name']):

            # Figure File Name
            img_fpath = os.path.join(self.image_dir, record['Figure file Name'])
            size = self._get_preferred_image_size(img_fpath)
            doc.add_picture(img_fpath, width=size[0])  # , height=size[1])
            p = doc.paragraphs[-1]
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

            # Figure Comment
            items = self._to_array(record['Figure comment'], '\n')
            first = True
            for item in items:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = docx.shared.Pt(10)
                p.paragraph_format.space_after = docx.shared.Pt(0)
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
                if first:
                    p.add_run('Figure: ').bold = True
                    first = False
                p.add_run(item)

        p.paragraph_format.space_after = docx.shared.Pt(14)

        # References
        items = self._to_array(record['References'], '\n')
        first = True
        for item in items:
            if first:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = docx.shared.Pt(11)
                p.paragraph_format.space_after = docx.shared.Pt(0)
                p.add_run('References:').bold = True
                first = False
            p = doc.add_paragraph()
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = docx.shared.Pt(10)
            p.paragraph_format.space_after = docx.shared.Pt(0)
            p.add_run(item)
        p.paragraph_format.space_after = docx.shared.Pt(10)

        # Acknowledgement
        items = self._to_array(record['Acknowledgement'], '\n')
        first = True
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = docx.shared.Pt(10)
            p.paragraph_format.space_after = docx.shared.Pt(0)
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            if first:
                p.add_run('Ackknowledgement: ').bold = True
                first = False
            p.add_run(item)
        p.paragraph_format.space_after = docx.shared.Pt(10)

        # Funding
        items = self._to_array(record['Funding'], '\n')
        first = True
        for item in items:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = docx.shared.Pt(10)
            p.paragraph_format.space_after = docx.shared.Pt(0)
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            if first:
                p.add_run('Funding: ').bold = True
                first = False
            p.add_run(item)
        p.paragraph_format.space_after = docx.shared.Pt(10)

        # Citation
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = docx.shared.Pt(10)
        p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run('Citation: ').bold = True
        author_tmp = ''
        first = True
        for author in authors:
            m = self.exreg4author.match(author)
            if not first:
                author_tmp += ', '
            author_tmp += m.group(1).strip()
            first = False
        p.add_run(author_tmp + ' (2016). ' + record['Title'].strip().replace('\n', ' ') + '. ')
        p.add_run('Advances in Neuroinformatics IV. ').italic = True
        p.add_run(
            'AINI 2016 and INCF Nodes Workshop Abstract: ' + record['Program No. Long'].strip() + '. DOI:' + record[
                'DOI'].strip())


if __name__ == '__main__':
    img_dir = './image'
    input_xlsx = 'input.xlsx'
    output_docx = 'output.docx'
    template_docx = './template/aini2016.docx'
